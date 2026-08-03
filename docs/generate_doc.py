#!/usr/bin/env python3
"""
Generates Free_Lector_Dokumentacija.docx
Professional project documentation in Serbian.
Uses proper Word formatting (no markdown artifacts, no raw dashes for bullets).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(3.0)
section.right_margin = Cm(3.0)
section.top_margin = Cm(2.4)
section.bottom_margin = Cm(2.4)

# Style configuration
style = doc.styles['Normal']
style.font.name = 'Cambria'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5


for level in range(1, 4):
    hstyle = doc.styles[f'Heading {level}']
    hstyle.font.name = 'Cambria'
    hstyle.font.bold = True
    if level == 1:
        hstyle.font.size = Pt(18)
    elif level == 2:
        hstyle.font.size = Pt(14)
    else:
        hstyle.font.size = Pt(12)


def add_bullet(text):
    """Add a proper Word bullet list item."""
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def add_numbered(text):
    """Add a numbered list item."""
    p = doc.add_paragraph(text, style='List Number')
    return p


def add_code_run(paragraph, text):
    """Add a code-formatted run within a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    return run



# === TITLE PAGE ===
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FREE LECTOR')
run.font.name = 'Cambria'
run.font.size = Pt(24)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Deterministička lektorska provera dokumenata')
run.font.name = 'Cambria'
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Rule-based pristup bez AI modela')
run.font.name = 'Cambria'
run.font.size = Pt(14)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Projektna dokumentacija')
run.font.name = 'Cambria'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Verzija 1.0 \u2013 Avgust 2026')
run.font.name = 'Cambria'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Autor: Tomislav Fedek')
run.font.name = 'Cambria'
run.font.size = Pt(12)

doc.add_page_break()



# === TABLE OF CONTENTS ===
doc.add_heading('Sadržaj', level=1)
toc_items = [
    '1. Uvod',
    '2. Cilj projekta',
    '3. Metodologija razvoja',
    '4. Arhitektura sistema',
    '5. Kategorije provera',
    '6. Evaluation framework',
    '7. Alati korišćeni u razvoju',
    '8. Ograničenja',
    '9. Zaključak',
]
for item in toc_items:
    doc.add_paragraph(item)
doc.add_page_break()



# === 1. UVOD ===
doc.add_heading('1. Uvod', level=1)
doc.add_paragraph(
    'Free Lector je alat za determinističku lektorsku proveru dokumenata koji radi '
    'isključivo u korisnikovom web pregledaču. Ne koristi veštačku inteligenciju, ne šalje '
    'podatke ni na jedan server, i ne zahteva instalaciju. Ceo proces parsiranja dokumenta, '
    'primene pravila i generisanja izveštaja odvija se lokalno, u jednom HTML fajlu sa '
    'pratećim JavaScript modulima.'
)
doc.add_paragraph(
    'Alat je namenjen autorima akademskih radova, lektorima i urednicima koji žele brzu, '
    'ponovljivu proveru tehničkih artefakata u tekstu pre slanja na recenziju ili '
    'objavljivanje. Free Lector ne zamenjuje ljudskog lektora \u2013 on hvata ono što ljudsko '
    'oko lako previdi: duple razmake, neuparene zagrade, mešanje pisama, ravne navodnike, '
    'markdown ostatke iz konverzija, i slične sistematske greške.'
)
doc.add_paragraph(
    'Projekat je open-source, hostovan na GitHub Pages, i dostupan svima besplatno na '
    'adresi: https://tfedek.github.io/free-lector/'
)



# === 2. CILJ ===
doc.add_heading('2. Cilj projekta', level=1)
doc.add_paragraph('Cilj projekta je višestruk:')
add_numbered(
    'Pružiti besplatan, pristupačan alat za tehničku proveru dokumenata na srpskom jeziku.'
)
add_numbered(
    'Demonstrirati da se značajan deo lektorskih provera može automatizovati '
    'determinističkim pravilima bez potrebe za AI modelima ili cloud servisima.'
)
add_numbered(
    'Podržati DOCX format (Microsoft Word) sa punim parsiranjem OOXML strukture, '
    'uključujući tabele, fusnote, endnote, headers/footers, tracked changes, i numerisane liste.'
)
add_numbered(
    'Omogućiti izvoz nalaza u Excel, Markdown i JSON format za dalju obradu.'
)
add_numbered(
    'Uspostaviti evaluation framework koji omogućava merenje tačnosti alata '
    'na različitim korpusima dokumenata.'
)
add_numbered(
    'Održati privatnost korisnika \u2013 nijedan bajt dokumenta ne napušta pregledač.'
)



# === 3. METODOLOGIJA ===
doc.add_heading('3. Metodologija razvoja', level=1)

doc.add_heading('3.1 Uloga AI alata u razvoju', level=2)
doc.add_paragraph(
    'Free Lector je razvijen u saradnji sa AI asistentom (Claude, Anthropic). '
    'Važno je razlikovati dve potpuno različite stvari:'
)
doc.add_paragraph(
    'Sam alat ne koristi AI. On je čisto rule-based: deterministička pravila, '
    'regex obrasci, parsiranje XML strukture. Isti ulaz uvek daje isti izlaz.'
)
doc.add_paragraph(
    'Razvoj alata je koristio AI asistenta kao programerski alat. AI je pisao kod, '
    'predlagao arhitekturu, generisao testove, i pomagao u debugging-u. Ovo je analogno '
    'korišćenju IDE-a sa intellisense-om ili Stack Overflow-a \u2013 alat za produktivnost '
    'developera, ne komponenta samog proizvoda.'
)
doc.add_paragraph('Konkretno, AI asistent je korišćen za:')
add_bullet('Pisanje inicijalnog koda parsera i rule engine-a na osnovu specifikacija')
add_bullet('Implementaciju naprednog OOXML parsiranja (gridSpan, vMerge, tracked changes, nested tables)')
add_bullet('Dizajn evaluation framework-a (Wilson CI, Brier score, reliability bins)')
add_bullet('Property-based fuzz testiranje sa fast-check bibliotekom')
add_bullet('Debugging false positive-a na realnim dokumentima')
add_bullet('Iterativno poboljšanje pravila kroz više rundi revizije')
add_bullet('Generisanje ove dokumentacije')
doc.add_paragraph(
    'Sav generisani kod je pregledan, testiran i modifikovan od strane autora. '
    'AI asistent nema pristup korisničkim dokumentima niti učestvuje u radu alata '
    'nakon što je kod napisan.'
)



doc.add_heading('3.2 Iterativni ciklus razvoja', level=2)
doc.add_paragraph('Razvoj je prolazio kroz više faza, svaka sa specifičnim ciljem:')

p = doc.add_paragraph()
run = p.add_run('Faza 1 \u2013 Osnovna funkcionalnost: ')
run.bold = True
p.add_run(
    'DOCX parsiranje, osnovne provere (razmaci, zagrade, navodnici, duple reči), '
    'HTML interfejs, Excel/Markdown/JSON izvoz.'
)

p = doc.add_paragraph()
run = p.add_run('Faza 2 \u2013 Napredno OOXML parsiranje: ')
run.bold = True
p.add_run(
    'Podrška za gridSpan, vMerge, tracked changes, nested tables, '
    'numerisanje sa lvlOverride/lvlRestart, header/footer parsiranje.'
)

p = doc.add_paragraph()
run = p.add_run('Faza 3 \u2013 Fuzz testiranje: ')
run.bold = True
p.add_run(
    'Property-based testovi sa fast-check bibliotekom. 10 svojstava, 500\u201310000 slučajeva. '
    'Otkriveni i popravljeni edge case-ovi.'
)

p = doc.add_paragraph()
run = p.add_run('Faza 4 \u2013 Testiranje na realnim dokumentima: ')
run.bold = True
p.add_run(
    'Provera na akademskim radovima (doktorska disertacija, klasičnofilološke studije, '
    'istorijski radovi). Eliminacija false positive-a: bibliografija, TOC, naslovna strana.'
)

p = doc.add_paragraph()
run = p.add_run('Faza 5 \u2013 Preset sistem i evaluation: ')
run.bold = True
p.add_run(
    'Basic/Full/Custom presets. Wilson score confidence intervali. Brier score. '
    'Cochran formula za veličinu uzorka. Stroga CSV validacija.'
)



doc.add_heading('3.3 Testiranje i validacija', level=2)
doc.add_paragraph('Projekat ima tri nivoa testiranja:')
add_numbered(
    'Unit testovi (109 testova) \u2013 pokrivaju svaki rule, svaki edge case parsera, '
    'svaku kombinaciju opcija.'
)
add_numbered(
    'Property-based fuzz testovi (10 svojstava, 500 slučajeva svako, 10000 za nightly) \u2013 '
    'generišu nasumične ulaze i proveravaju invarijante: nema crash-a, nema gubitka teksta, '
    'determinizam, konzistentnost ID-jeva.'
)
add_numbered(
    'Evaluation testovi (30 testova) \u2013 proveravaju CLI alate za evaluaciju: '
    'validaciju CSV-a, Wilson interval izračunavanje, sample size formulu, '
    'odbijanje nevalidnih ulaza.'
)
doc.add_paragraph('Ukupno: 149 automatskih provera koje se pokreću pre svakog commit-a.')



# === 4. ARHITEKTURA ===
doc.add_heading('4. Arhitektura sistema', level=1)
doc.add_paragraph(
    'Sistem se sastoji od pet glavnih modula koji se učitavaju u pregledač kao '
    'obični JavaScript fajlovi (bez build step-a, bez bundler-a):'
)

doc.add_heading('4.1 Parser (parser.js)', level=2)
doc.add_paragraph(
    'Parsira DOCX (ZIP arhivu sa XML fajlovima) koristeći JSZip biblioteku. '
    'Izvlači strukturirane podatke iz OOXML formata:'
)
add_bullet('Paragrafi sa stilovima, run-ovima, i formatiranjem')
add_bullet('Tabele sa gridSpan (spojene kolone) i vMerge (spojeni redovi)')
add_bullet('Fusnote i endnote')
add_bullet('Headers i footers')
add_bullet('Tracked changes \u2013 detektuje, ne primenjuje')
add_bullet('Numerisane liste sa lvlOverride i lvlRestart')
add_bullet('Stilovi za detekciju naslova i TOC unosa')
add_bullet('ZIP bezbednost: odbijanje ekstreman compression ratio i zip-bomb detekcija')
doc.add_paragraph(
    'Parser proizvodi docMap objekat koji sadrži sve elemente dokumenta sa metapodacima '
    'potrebnim za rule engine (indeksi, sekcije, tip elementa, runs za per-character analizu).'
)



doc.add_heading('4.2 Rule Engine (rules.js)', level=2)
doc.add_paragraph(
    'Deterministički engine koji primenjuje pravila na parsiranu strukturu dokumenta. '
    'Svako pravilo je obična JavaScript funkcija koja prima docMap i vraća niz nalaza. '
    'Nema ML modela, nema heuristika koje se menjaju \u2013 isti ulaz uvek daje isti izlaz.'
)
doc.add_paragraph('Svaki nalaz (finding) sadrži:')
add_bullet('rule_id \u2013 stabilan identifikator pravila (npr. spacing_body, brackets_header)')
add_bullet('category \u2013 kategorija greške (Razmaci, Zagrade, Tipografija...)')
add_bullet('priority \u2013 Obavezno, Blokirajuće, Proveriti, ili Preporuka')
add_bullet('confidence \u2013 heuristička procena pouzdanosti (0.0\u20131.0)')
add_bullet('original \u2013 originalni tekst sa greškom')
add_bullet('replacement \u2013 predlog ispravke')
add_bullet('rationale \u2013 obrazloženje zašto je ovo greška')
add_bullet('section \u2013 lokacija u dokumentu')
doc.add_paragraph(
    'Napomena: confidence je heuristički rang, ne empirijski kalibrisan '
    'verovatnosni skor. Služi za sortiranje nalaza po ozbiljnosti.'
)

doc.add_heading('4.3 Preset sistem (presets.js)', level=2)
doc.add_paragraph('Dva ugrađena preset-a određuju koje provere su aktivne:')
p = doc.add_paragraph()
run = p.add_run('Osnovna provera (Basic) ')
run.bold = True
p.add_run(
    '\u2013 samo visoko-pouzdane provere koje retko daju false positive: razmaci, zagrade, '
    'navodnici, mešanje pisama, duple reči, fusnote. Headers/footers su isključeni. '
    'Ovo je podrazumevani režim.'
)
p = doc.add_paragraph()
run = p.add_run('Puna provera (Full) ')
run.bold = True
p.add_run(
    '\u2013 sve provere uključene: bibliografija, TOC, numeracija, URL-ovi, grčki citati, '
    'markdown artefakti, ALL-CAPS, headers/footers. '
    'Može dati više false positive-a ali ništa ne propušta.'
)
p = doc.add_paragraph()
run = p.add_run('Prilagođeno (Custom) ')
run.bold = True
p.add_run(
    '\u2013 korisnik ručno bira koje provere želi. Automatski se aktivira '
    'kada korisnik promeni bilo koji checkbox iz Basic ili Full režima.'
)



doc.add_heading('4.4 Eksporter (exporter.js)', level=2)
doc.add_paragraph('Generiše tri formata izveštaja:')
add_bullet(
    'Excel (.xlsx) \u2013 koristi SheetJS biblioteku. Sadrži sve nalaze sa kolonama: '
    'lokacija, kategorija, prioritet, original, predlog, obrazloženje, pouzdanost, rule_id, preset.'
)
add_bullet(
    'Markdown (.md) \u2013 tabelarni format pogodan za pregled u bilo kom tekst editoru.'
)
add_bullet(
    'JSON (.json) \u2013 mašinsko-čitljiv format za dalju obradu ili integraciju sa drugim alatima.'
)

doc.add_heading('4.5 Korisnički interfejs (app.js, index.html)', level=2)
doc.add_paragraph(
    'Jednostrana web aplikacija bez framework-a. Čist HTML/CSS/JS. '
    'Korisnik upload-uje DOCX fajl, bira preset, pokreće audit, i dobija tabelarni '
    'prikaz nalaza sa filterima po prioritetu, kategoriji i sekciji. '
    'Svaki nalaz ima dugmiće za prihvatanje ili odbijanje.'
)
doc.add_paragraph(
    'Interfejs prikazuje i sumarni status: koliko obaveznih, blokirajućih, i preporučenih '
    'nalaza postoji, i da li je dokument spreman za objavljivanje (0 obaveznih i 0 blokirajućih).'
)



# === 5. KATEGORIJE PROVERA ===
doc.add_heading('5. Kategorije provera', level=1)
doc.add_paragraph('Alat proverava sledeće kategorije tehničkih grešaka:')

checks = [
    ('5.1 Razmaci i interpunkcija',
     'Detektuje duple razmake, razmak pre zapete/tačke/tačke-zareza, razmak pre zatvorene '
     'zagrade, razmak posle otvorene zagrade, razmak na početku reda. '
     'Isključuje višestruke razmake u code blokovima i tabelarnim ćelijama.'),
    ('5.2 Nebalansirane zagrade',
     'Proverava da li svaka otvorena zagrada ima odgovarajuću zatvorenu. '
     'Detektuje i tačku-zarez unutar nezatvorene zagrade (čest artefakt iz find-replace operacija).'),
    ('5.3 Tipografski navodnici',
     'Na srpskom jeziku standard su navodnici \u201e...\u201c (donji i gornji). Alat detektuje '
     'ravne navodnike i neuparene tipografske navodnike.'),
    ('5.4 Mešanje ćirilice i latinice',
     'Detektuje reči koje sadrže i ćirilične i latinične karaktere (npr. latinično \u201ea\u201c '
     'umesto ćirilicnog \u201ea\u201c). Isključuje poznate skraćenice i mešovite kontekste.'),
    ('5.5 Duple reči',
     'Detektuje uzastopno ponovljene reči (\u201ei i\u201c, \u201ena na\u201c, \u201eje je\u201c). '
     'Isključuje namerne duplikate u citatima i specijalnim kontekstima.'),
    ('5.6 ALL-CAPS u tekstu',
     'Prijavljuje reči napisane velikim slovima u telu teksta koje nisu poznate '
     'skraćenice. Isključuje TOC unose, naslove, tabelarne ćelije, naslovnu stranu.'),
    ('5.7 Markdown artefakti',
     'Detektuje ostatke markdown formatiranja u Word dokumentu: vidljive zvezdice, '
     'hash naslove, bullet tačkice, backtick blokove.'),
    ('5.8 Bibliografija',
     'Proverava da li bibliografski unosi imaju godinu izdanja. Prepoznaje format '
     'akademskih referenci i ne prijavljuje antičke izvore ili časopisne članke.'),
    ('5.9 TOC vs naslovi',
     'Poredi sadržaj (Table of Contents) sa stvarnim naslovima u telu dokumenta. '
     'Prijavljuje neslaganja u tekstu ili redosledu.'),
    ('5.10 URL validacija',
     'Detektuje URL-ove koji se završavaju interpunkcijskim znakom (tačka, zarez) '
     'što može značiti da je URL skraćen ili nevalidan.'),
    ('5.11 Grčki citati bez prevoda',
     'Za akademske radove: detektuje grčke citate koji nemaju prevod u neposrednom '
     'okruženju. Isključeno u Basic modu (niche provera).'),
    ('5.12 Fusnote i endnote',
     'Proverava prazne fusnote/endnote. U punom modu proverava i sadržaj fusnota '
     'za razmake, zagrade, mešanje pisama.'),
    ('5.13 Headers i footers',
     'Kada je uključeno (Full mod), primenjuje sve aktivne provere i na sadržaj '
     'header-a i footer-a dokumenta.'),
]
for title, desc in checks:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)



# === 6. EVALUATION FRAMEWORK ===
doc.add_heading('6. Evaluation framework', level=1)
doc.add_paragraph(
    'Za merenje tačnosti alata uspostavljen je evaluation framework koji omogućava '
    'sistematsko poređenje nalaza sa ljudskim oznakama (labels).'
)

doc.add_heading('6.1 Wilson score interval', level=2)
doc.add_paragraph(
    'Za svako pravilo (rule_id), izračunava se Wilson score confidence interval koji daje '
    'donju i gornju granicu tačnosti sa 95% pouzdanosti. Wilson interval je bolji od '
    'normalnog intervala za male uzorke i proporcije blizu 0 ili 1.'
)
doc.add_paragraph(
    'Primer: ako pravilo spacing_body ima 45 od 50 tačnih nalaza, Wilson interval '
    'može biti [0.78, 0.96], što znači da smo 95% sigurni da je prava tačnost '
    'između 78% i 96%.'
)

doc.add_heading('6.2 Brier score', level=2)
doc.add_paragraph(
    'Meri kalibraciju confidence procena. Brier score je srednja kvadratna greška '
    'između predicted confidence i stvarnog ishoda (0 ili 1). Idealan skor je 0.0, '
    'a skor ispod 0.25 se smatra dobrom kalibracijom.'
)

doc.add_heading('6.3 Reliability bins', level=2)
doc.add_paragraph(
    'Grupisanje nalaza po confidence opsezima (0.0\u20130.2, 0.2\u20130.4, ... 0.8\u20131.0) '
    'i poređenje srednje predikcije sa stvarnom tačnošću u svakom binu. '
    'Vizualizuje kalibracioni dijagram \u2013 koliko se predviđanja poklapaju sa stvarnošću.'
)

doc.add_heading('6.4 Sample size (Cochran formula)', level=2)
doc.add_paragraph(
    'Cochran-ova formula određuje minimalni broj obeleženih nalaza potreban za '
    'statistički validnu procenu. Podržava design effect za korelaciju unutar '
    'istog dokumenta (nalazi u jednom dokumentu nisu nezavisni).'
)
doc.add_paragraph(
    'Primer: za marginu greške 5% i confidence 95%, sa p=0.5 i design effect 1, '
    'potrebno je oko 385 obeleženih nalaza.'
)



# === 7. ALATI ===
doc.add_heading('7. Alati korišćeni u razvoju', level=1)
doc.add_paragraph('Sledeći alati i tehnologije su korišćeni tokom razvoja projekta:')

tools = [
    ('Programski jezik', 'JavaScript (ES2020+), bez TypeScript-a, bez build step-a'),
    ('AI asistent', 'Claude (Anthropic) \u2013 za generisanje koda, debugging, testiranje'),
    ('IDE/Editor', 'Kiro CLI \u2013 AI-powered terminal razvojno okruženje'),
    ('Verziona kontrola', 'Git, GitHub (https://github.com/tfedek/free-lector)'),
    ('Hosting', 'GitHub Pages \u2013 statičko hostovanje bez servera'),
    ('Testiranje', 'Custom test framework (test.js), fast-check za property-based testove'),
    ('DOCX parsiranje', 'JSZip (ZIP ekstrakcija), ručno XML parsiranje (DOMParser)'),
    ('Excel izvoz', 'SheetJS (xlsx.full.min.js)'),
    ('CSS', 'Čist CSS bez framework-a, responsive dizajn'),
    ('Dokumentacija', 'python-docx za generisanje ovog dokumenta'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Komponenta'
hdr_cells[1].text = 'Tehnologija'
for comp, tech in tools:
    row_cells = table.add_row().cells
    row_cells[0].text = comp
    row_cells[1].text = tech

doc.add_paragraph()
doc.add_paragraph(
    'Napomena o transparentnosti: Ovaj projekat je razvijen uz pomoć AI asistenta. '
    'To je eksplicitno dokumentovano jer verujemo da je transparentnost važna. '
    'AI je korišćen kao alat za produktivnost (slično korišćenju IDE autocomplete-a '
    'ili kopiranja koda sa StackOverflow-a), ne kao zamena za razumevanje koda. '
    'Sav kod je pregledan, testiran, i modifikovan od strane autora.'
)



# === 8. OGRANICENJA ===
doc.add_heading('8. Ograničenja', level=1)
doc.add_paragraph(
    'Free Lector ima jasna ograničenja koja proizlaze iz rule-based pristupa:'
)
add_numbered(
    'Ne razume kontekst \u2013 ne može oceniti da li je rečenica logički ispravna, '
    'samo da li ima tehničke artefakte.'
)
add_numbered(
    'Ne proverava gramatiku \u2013 padežne greške, slaganje subjekta i predikata, '
    'red reči \u2013 to zahteva NLP ili AI model.'
)
add_numbered(
    'Ne proverava stil \u2013 ne može reći da li je tekst preformalan, preopširan, '
    'ili neodgovarajuć za ciljnu publiku.'
)
add_numbered(
    'Ne proverava sadržajnu tačnost \u2013 ne može verifikovati da li su podaci, '
    'datumi, ili reference tačni.'
)
add_numbered(
    'Ne proverava prevode \u2013 ne može oceniti kvalitet prevoda grčkih ili '
    'latinskih citata.'
)
add_numbered(
    'Može dati false positive za neuobičajene ali validne konstrukcije \u2013 '
    'zato svaki nalaz ima confidence procenu i prioritet.'
)
add_numbered(
    'Podržava samo DOCX format. PDF, ODT, i drugi formati nisu podržani.'
)
add_numbered(
    'Nema OCR \u2013 ne može čitati skenirane dokumente.'
)
doc.add_paragraph(
    'Free Lector je alat koji dopunjuje, ne zamenjuje, rad ljudskog lektora. '
    'On hvata tehničke artefakte koje ljudsko oko lako previdi pri dugom čitanju, '
    'ali ne može doneti sud o kvalitetu teksta.'
)



# === 9. ZAKLJUCAK ===
doc.add_heading('9. Zaključak', level=1)
doc.add_paragraph(
    'Free Lector demonstrira da je moguće napraviti koristan, besplatan alat za '
    'tehničku lektorsku proveru dokumenata bez AI modela i bez slanja podataka na server. '
    'Deterministički, rule-based pristup ima prednosti u transparentnosti '
    '(korisnik tačno zna šta se proverava), ponovljivosti (isti ulaz = isti izlaz), '
    'i privatnosti (ništa ne napušta pregledač).'
)
doc.add_paragraph(
    'Projekat je testiran na realnim akademskim dokumentima (doktorske disertacije, '
    'klasičnofilološke studije, istorijski radovi) i pokazao je da može uhvatiti '
    'tehničke greške koje bi inače prošle neprimećene, bez generisanja prevelikog '
    'broja lažnih alarma.'
)
doc.add_paragraph(
    'Korišćenje AI asistenta u razvoju je transparentno dokumentovano. Sam alat '
    'ne sadrži AI komponentu \u2013 to su čista, deterministička pravila koja se mogu '
    'pregledati, razumeti, i modifikovati.'
)
doc.add_paragraph(
    'Izvorni kod je javno dostupan na: https://github.com/tfedek/free-lector'
)
doc.add_paragraph(
    'Alat je dostupan na: https://tfedek.github.io/free-lector/'
)


# === SAVE ===
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Free_Lector_Dokumentacija.docx')
doc.save(output_path)
print(f'Dokument generisan: {output_path}')
print(f'Veličina: {os.path.getsize(output_path) / 1024:.1f} KB')
